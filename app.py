
import os
from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file
from create_model import db, Protocol, Teams, Referees, Authorize, Players
from datetime import datetime
from docx import Document
from functools import wraps
from werkzeug.utils import secure_filename
from pathlib import Path
from io import BytesIO

app = Flask(__name__)  # Исправлено: Передаем __name__
app.secret_key = os.environ.get('SECRET_KEY', 'your_default_secret_key')
app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://postgres:1703@localhost/KFV'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)

# Декоратор для проверки аутентификации
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Требуется вход в систему.', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/')
def index():
    return render_template('index1.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        fio = request.form.get('fio')
        password = request.form.get('password')
        user = Authorize.query.filter_by(fio=fio, password=password).first()
        if user:
            session['user_id'] = user.id
            return redirect(url_for('add_protocol'))  # Перенаправляем на страницу с формой
        else:
            flash('Неверный логин или пароль.')
    users = Authorize.query.all()
    return render_template('login.html', users=users)

@app.route('/logout')
def logout():
    session.pop('user_id', None)
    return redirect(url_for('login'))

def replace_all_text_in_document(document, protocol):
    """Заменяет текст в документе Word, включая таблицы, на основе данных из объекта Protocol."""
    mappings = {
        '{{competition_name}}': protocol.competition_name,
        '{{city}}': protocol.city,
        '{{hall}}': protocol.hall,
        '{{stage}}': protocol.stage,
        '{{match}}': str(protocol.match),
        '{{gender}}': protocol.gender,
        '{{category}}': protocol.category,
        '{{winner}}': str(protocol.winner),  # Use ID
        '{{link_protocol}}': protocol.link_protocol,
        '{{date}}': protocol.date.strftime("%Y-%m-%d"),
        '{{time}}': protocol.time.strftime("%H:%M"),
        '{{team1}}': str(protocol.team1),  # Use ID
        '{{team2}}': str(protocol.team2),  # Use ID
        '{{secretary}}': str(protocol.secretary),  # Use ID
        '{{referee}}': str(protocol.referee)  # Use ID
    }

    # Замена в параграфах (как было)
    for paragraph in document.paragraphs:
        for search_text, replace_text in mappings.items():
            if search_text in paragraph.text:
                paragraph.text = paragraph.text.replace(search_text, str(replace_text))

    # Замена в таблицах
    for i, table in enumerate(document.tables):
        print(f"Processing Table {i}: {len(table.rows)} rows, {len(table.columns)} columns") # Debug

        for j, row in enumerate(table.rows):
            for k, cell in enumerate(row.cells):
                print(f"  Cell ({j}, {k}): {cell.text}") # Debug
                for search_text, replace_text in mappings.items():
                    if search_text in cell.text:
                        cell.text = cell.text.replace(search_text, str(replace_text))
                        print(f"    Replaced '{search_text}' with '{replace_text}' in cell ({j}, {k})") #Debug

@app.route('/add_protocol', methods=['GET', 'POST'])
@login_required
def add_protocol():
    teams = Teams.query.all()
    referees = Referees.query.all()
    if request.method == 'POST':
        try:
            competition_name = request.form['competition_name']
            city = request.form['city']
            hall = request.form['hall']
            stage = request.form['stage']
            match = int(request.form['match'])
            gender = request.form['gender']
            category = request.form['category']
            winner = int(request.form['winner'])
            link_protocol = request.form['link_protocol']
            date_str = request.form['date']
            time_str = request.form['time']
            team1 = int(request.form['team1'])
            team2 = int(request.form['team2'])
            secretary = int(request.form['secretary'])
            referee = int(request.form['referee'])

            combined_datetime = datetime.strptime(f"{date_str} {time_str}", "%Y-%m-%d %H:%M")

            new_protocol = Protocol(
                competition_name=competition_name,
                city=city,
                hall=hall,
                stage=stage,
                match=match,
                gender=gender,
                category=category,
                winner=winner,
                link_protocol=link_protocol,
                date=combined_datetime,
                time=combined_datetime.time(),
                team1=team1,
                team2=team2,
                secretary=secretary,
                referee=referee
            )
            db.session.add(new_protocol)
            db.session.commit()

            # After saving, generate the document
            template_path = os.path.abspath(os.path.join('generated_docs', 'Shab.docx'))
            document = Document(template_path)
            replace_all_text_in_document(document, new_protocol)

            # Save the modified document
            output_path = f"protocol_{new_protocol.id}.docx"
            document.save(output_path)

            # Get the path of the selected file
            link_protocol_path = request.form['link_protocol']

            # Get the directory of the selected file
            link_protocol_dir = os.path.dirname(link_protocol_path)

            # Save the document in the selected directory
            document.save(os.path.join(link_protocol_dir, f"protocol_{new_protocol.id}.docx"))

            flash('Протокол успешно добавлен и сгенерирован!', 'success')
            return redirect(url_for('download_protocol', filename=os.path.join(link_protocol_dir, f"protocol_{new_protocol.id}.docx"))) # Redirect to download

        except Exception as e:
            db.session.rollback()
            flash(f'Ошибка добавления протокола: {e}', 'error')
            return render_template('protocols.html', error=str(e), teams=teams, referees=referees)

    # Загрузка данных для выпадающих списков
    return render_template('protocols.html', teams=teams, referees=referees)

@app.route('/download_protocol/<filename>')
def download_protocol(filename):
    return send_file(filename, as_attachment=True)


if __name__ == '__main__':
     with app.app_context():
        db.create_all()
     app.run(debug=True)
